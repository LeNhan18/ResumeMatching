import os
import sys
import logging
from RAG.pipeline import CVMatcherPipeline
from RAG.schemas import ScoringWeights

# Set up logging to print to stdout
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger("CVMatcherDemo")

def create_sample_docx_cv(file_path: str):
    """Generates a mock CV Word document using python-docx to verify parser compatibility."""
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("RESUME: NGUYEN PHAN TIEN", 0)
        
        # Personal info
        p = doc.add_paragraph()
        p.add_run("Email: tien.nguyen@example.com | Phone: 0987654321 | Location: Hanoi, Vietnam\n")
        p.add_run("Industry: IT | Sector: Fintech / Backend Engineering")
        
        # Summary
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(
            "Experienced Backend Engineer with over 4 years of hands-on expertise building scalable "
            "financial transaction web APIs, managing PostgreSQL databases, and deploying containerized "
            "applications. Proficient in Python, FastAPI, Docker, and Git workflow."
        )
        
        # Skills
        doc.add_heading("Core Technical Skills", level=1)
        doc.add_paragraph("Python, FastAPI, Docker, SQL, PostgreSQL, Git, Redis, RESTful APIs, Linux")
        
        # Experience
        doc.add_heading("Work Experience", level=1)
        
        # Job 1
        p1 = doc.add_paragraph()
        p1.add_run("1. Senior Python Developer | Fintech Solutions Corp\n").bold = True
        p1.add_run("Duration: 2022-03 to Present (Ongoing)\n")
        p1.add_run("Responsibilities & Achievements:\n")
        doc.add_paragraph(
            "- Architected and optimized payment gateway microservices using Python and FastAPI, handling 10k+ requests/min.\n"
            "- Managed and query-tuned large-scale PostgreSQL databases.\n"
            "- Streamlined deployment pipeline by Dockerizing service instances.\n"
            "Skills Used: Python, FastAPI, Docker, PostgreSQL, Redis"
        )
        
        # Job 2
        p2 = doc.add_paragraph()
        p2.add_run("2. Junior Backend Developer | Tech Startup Lab\n").bold = True
        p2.add_run("Duration: 2020-01 to 2022-02 (25 months)\n")
        p2.add_run("Responsibilities & Achievements:\n")
        doc.add_paragraph(
            "- Developed and maintained internal ERP system APIs using Flask and SQLite.\n"
            "- Collaborated with frontend developers to integrate client portals.\n"
            "Skills Used: Python, Flask, SQLite, Git"
        )
        
        # Education
        doc.add_heading("Education", level=1)
        doc.add_paragraph("Bachelor of Software Engineering | Hanoi University of Science and Technology (HUST)\nGraduation Year: 2020")
        
        # Certifications
        doc.add_heading("Certifications", level=1)
        doc.add_paragraph("- AWS Certified Developer - Associate\n- Docker Certified Associate")
        
        doc.save(file_path)
        logger.info(f"Sample DOCX CV generated successfully at: {file_path}")
    except ImportError:
        logger.warning("python-docx is not installed in the active environment. Generating text fallback instead.")
        create_sample_text_cv(file_path.replace(".docx", ".txt"))

def create_sample_text_cv(file_path: str):
    """Generates a mock CV text file."""
    content = """
    RESUME: TRAN HONG NHUNG
    Email: nhung.tran@example.com
    Phone: 0912345678
    Location: Ho Chi Minh City, Vietnam
    Industry: IT
    
    SUMMARY:
    Highly motivated Software Engineer specialized in React and Frontend development. 3 years of experience.
    Strong understanding of UI/UX layouts, state management, and modern javascript.
    
    CORE SKILLS:
    ReactJS, JavaScript, HTML5, CSS3, TailwindCSS, Git, REST APIs
    
    WORK EXPERIENCE:
    1. Frontend Developer | E-Commerce Hub
       Position: Frontend Developer
       Duration: 2021-06 to Present (Ongoing)
       Skills used: ReactJS, JavaScript, TailwindCSS, Git
       Description: Developed custom storefront components, optimized page loading speeds by 30%.
       
    2. Junior Developer | Web agency
       Position: Web developer
       Duration: 2020-03 to 2021-05
       Skills used: HTML, CSS, JavaScript, Git
       Description: Built responsive website landing pages.
       
    EDUCATION:
    Bachelor of Information Technology | HCMC University of Science (HCMUS)
    Graduation: 2020
    """
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    logger.info(f"Sample TXT CV generated successfully at: {file_path}")

def create_sample_jd(file_path: str):
    """Generates a mock Job Description text file."""
    content = """
    JOB DESCRIPTION: SENIOR BACKEND ENGINEER
    Company: Cyber Finance Inc.
    Location: Hanoi (Hybrid)
    Industry: IT
    Domain: Backend
    Level: Senior
    
    ROLE OVERVIEW:
    We are looking for a Senior Backend Engineer to join our Core Fintech team. You will design secure transaction systems, 
    optimize heavy SQL queries, and manage cloud infrastructure.
    
    REQUIREMENTS (MANDATORY):
    - Minimum 3.0 years of relevant professional backend experience.
    - Strong programming proficiency in Python.
    - Deep knowledge of containerization using Docker.
    - Practical experience with databases, particularly PostgreSQL.
    
    NICE TO HAVE (PREFERRED):
    - Experience with Kubernetes.
    - Knowledge of cloud architecture (AWS or GCP).
    - Familiarity with FastAPI or Go.
    
    EDUCATION:
    - Bachelor's degree in Computer Science, Software Engineering or related fields.
    """
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    logger.info(f"Sample JD generated successfully at: {file_path}")

def main():
    logger.info("=== STARTING RAG CV MATCHER PIPELINE DEMONSTRATION ===")
    
    # 1. Define sample file paths
    docx_cv_path = "sample_cv_1.docx"
    txt_cv_path = "sample_cv_2.txt"
    jd_path = "sample_jd.txt"
    
    # 2. Create sample files
    create_sample_docx_cv(docx_cv_path)
    create_sample_text_cv(txt_cv_path)
    create_sample_jd(jd_path)
    
    # 3. Initialize pipeline
    pipeline = CVMatcherPipeline()
    
    # 4. Ingest CVs
    logger.info("--- INGESTING CANDIDATE CVS ---")
    try:
        # Ingest DOCX CV (if docx cv wasn't generated due to import error, use the txt fallback instead)
        target_docx_path = docx_cv_path if os.path.exists(docx_cv_path) else "sample_cv_1.txt"
        if not os.path.exists(target_docx_path):
            create_sample_text_cv(target_docx_path)
            
        cv1 = pipeline.ingest_cv(target_docx_path, cv_id="cand_tien")
        logger.info(f"Ingested Candidate 1: {cv1.name} | Skills: {cv1.skills}")
        
        cv2 = pipeline.ingest_cv(txt_cv_path, cv_id="cand_nhung")
        logger.info(f"Ingested Candidate 2: {cv2.name} | Skills: {cv2.skills}")
    except Exception as e:
        logger.error(f"Failed to ingest CVs: {e}")
        return
        
    # 5. Run matching and ranking for the JD
    logger.info("--- RUNNING PIPELINE MATCHING AND RANKING ---")
    custom_weights = ScoringWeights(
        skills=0.45,       # 45% weight on technical skills
        experience=0.35,   # 35% weight on duration & relevance
        education=0.10,    # 10% weight on academic degrees
        culture_fit=0.10   # 10% weight on soft skills/vibe
    )
    
    try:
        ranked_candidates = pipeline.rank_candidates(
            jd_text_or_file=jd_path,
            weights=custom_weights,
            limit=5
        )
        
        # 6. Display results
        logger.info("=== FINAL RANKED EVALUATION RESULTS ===")
        for rank, match in enumerate(ranked_candidates):
            print(f"\n{rank + 1}. Candidate: {match['candidate_name']} (ID: {match['candidate_id']})")
            print(f"   OVERALL MATCH SCORE: {match['score']}%")
            print("   Score Breakdown:")
            print(f"      - Technical Skills Alignment: {match['breakdown']['skills']}/100 (Weight: {custom_weights.skills*100:.0f}%)")
            print(f"      - Experience & Domain Tenure: {match['breakdown']['experience']}/100 (Weight: {custom_weights.experience*100:.0f}%)")
            print(f"      - Education Credentials:    {match['breakdown']['education']}/100 (Weight: {custom_weights.education*100:.0f}%)")
            print(f"      - Culture Fit / Soft Skills: {match['breakdown']['culture_fit']}/100 (Weight: {custom_weights.culture_fit*100:.0f}%)")
            
            hard_match_info = match['hard_match']
            print("   Hard-Match Constraints:")
            print(f"      - Experience Target:        {'PASS' if hard_match_info['exp_passed'] else 'FAIL'} ({hard_match_info['relevant_years']} yrs calculated, requested {hard_match_info['required_years']} yrs)")
            print(f"      - Mandatory Skills Target:  {'PASS' if hard_match_info['skills_passed'] else 'FAIL'} (Found: {', '.join(hard_match_info['matched_required_skills'])})")
            print(f"      - Education Level Target:   {'PASS' if hard_match_info['education_passed'] else 'FAIL'}")
            print(f"      - Overall Hard Match:       {'PASSED' if hard_match_info['overall_hard_match_passed'] else 'FAILED'}")
            
            print(f"   Key Strengths: {', '.join(match['strengths'])}")
            print(f"   Missing/Weak Skills: {', '.join(match['missing_skills']) or 'None'}")
            print(f"   Qualitative Reasoning:\n      {match['reasoning']}")
            print("-" * 80)
            
    except Exception as e:
        logger.error(f"Failed to rank candidates: {e}", exc_info=True)

    # 7. Cleanup generated files
    logger.info("--- CLEANING UP GENERATED DEMO FILES ---")
    for file in [docx_cv_path, txt_cv_path, jd_path, "sample_cv_1.txt"]:
        if os.path.exists(file):
            try:
                os.remove(file)
                logger.info(f"Removed temporary file: {file}")
            except Exception as e:
                logger.warning(f"Could not remove temporary file '{file}': {e}")

if __name__ == "__main__":
    main()
