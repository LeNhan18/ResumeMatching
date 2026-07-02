import os
import logging
import fitz
from PIL import Image
import io
import base64
from LLM.client import LLMClient
logger = logging.getLogger(__name__)

def parse_pdf_native(file_path: str) -> str:
    """Extracts text from a native PDF using pdfplumber, falling back to PyMuPDF."""
    text_content = []
    
    # Try pdfplumber first
    try:
        import pdfplumber
        logger.info(f"Attempting to parse native PDF with pdfplumber: {file_path}")
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(layout=True)  # maintain visual structure
                if page_text:
                    text_content.append(page_text)
        
        extracted = "\n".join(text_content).strip()
        if extracted:
            return extracted
    except ImportError:
        logger.warning("pdfplumber is not installed. Falling back to PyMuPDF (fitz).")
    except Exception as e:
        logger.error(f"Error parsing PDF with pdfplumber: {e}")

    # Fallback to PyMuPDF
    try:
        import fitz  # PyMuPDF
        logger.info(f"Attempting to parse native PDF with PyMuPDF: {file_path}")
        doc = fitz.open(file_path)
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_content.append(page_text)
        return "\n".join(text_content).strip()
    except ImportError:
        logger.warning("fitz (PyMuPDF) is not installed.")
    except Exception as e:
        logger.error(f"Error parsing PDF with PyMuPDF: {e}")
        
    return ""

def parse_docx(file_path: str) -> str:
    """Extracts text from a DOCX document using python-docx."""
    try:
        import docx
        logger.info(f"Parsing DOCX with python-docx: {file_path}")
        doc = docx.Document(file_path)
        
        # Read paragraphs
        paragraphs = [p.text for p in doc.paragraphs]
        
        # Read tables to capture structured experiences
        tables_content = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                tables_content.append(" | ".join(filter(None, row_text)))
                
        return "\n".join(paragraphs + tables_content).strip()
    except ImportError:
        logger.error("python-docx is not installed.")
        return ""
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return ""

def pdf_to_images(pdf_path: str) -> list:
    """Renders all pages of a PDF into PIL Images using PyMuPDF."""

    images = []
    try:
        logger.info(f"Rendering PDF pages to images for OCR: {pdf_path}")
        doc = fitz.open(pdf_path)
        for page in doc:
            pix = page.get_pixmap(dpi=150)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            images.append(img)
    except Exception as e:
        logger.error(f"Error rendering PDF pages to images: {e}")
    return images

def parse_with_surya(images: list) -> str:
    """Performs local OCR using Surya OCR package."""
    try:
        from surya.ocr import run_ocr
        from surya.model.detection import load_model as load_det_model, load_processor as load_det_processor
        from surya.model.recognition import load_model as load_rec_model, load_processor as load_rec_processor
        
        logger.info("Initializing Surya OCR models...")
        langs = ["vi", "en"]
        
        det_model = load_det_model()
        det_processor = load_det_processor()
        rec_model = load_rec_model()
        rec_processor = load_rec_processor()
        
        logger.info("Running Surya OCR on images...")
        predictions = run_ocr(images, [langs] * len(images), det_model, det_processor, rec_model, rec_processor)
        
        text_content = []
        for pred in predictions:
            page_text = "\n".join([line.text for line in pred.text_lines])
            text_content.append(page_text)
            
        return "\n\n".join(text_content)
    except ImportError:
        logger.warning("surya package is not installed. Cannot run Surya OCR.")
        raise
    except Exception as e:
        logger.error(f"Error during Surya OCR processing: {e}")
        raise

def parse_with_vlm(images: list) -> str:
    """Performs cloud-based OCR/transcription using a Vision-Language Model via OpenRouter."""

    llm = LLMClient()
    if not llm.is_configured():
        logger.warning("LLM client is not configured. Cannot call VLM API.")
        return ""
        
    vlm_model = os.getenv("OPENROUTER_VLM_MODEL", "qwen/qwen3-vl-32b-instruct")
    text_results = []
    
    for idx, img in enumerate(images):
        try:
            logger.info(f"Sending page {idx+1}/{len(images)} to VLM OCR model: {vlm_model}")
            buffered = io.BytesIO()
            img.save(buffered, format="JPEG")
            img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
            
            prompt = "Extract and transcribe all text from this page. Keep the layout and visual structure (e.g. columns, lists, experience blocks)."
            
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{img_base64}"
                            }
                        }
                    ]
                }
            ]
            
            response = llm.client.chat.completions.create(
                model=vlm_model,
                messages=messages,
                temperature=0.1
            )
            text_results.append(response.choices[0].message.content.strip())
        except Exception as e:
            logger.error(f"Error during VLM OCR on page {idx+1}: {e}")
            
    return "\n\n".join(text_results)

def parse_image_ocr(file_path: str) -> str:
    """
    Renders scanned PDFs or opens image files, then performs OCR using Surya (local) or VLM (API fallback).
    """
    logger.info(f"OCR/VLM parsing initiated for: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    
    # 1. Render/load images
    images = []
    if ext == ".pdf":
        images = pdf_to_images(file_path)
    elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
        try:
            images = [Image.open(file_path)]
        except Exception as e:
            logger.error(f"Error opening image file '{file_path}': {e}")
            return ""
            
    if not images:
        logger.error(f"No pages/images could be extracted for OCR from: {file_path}")
        return ""
        
    # 2. Decide OCR Engine based on config and availability
    do_ocr = os.getenv("DO_OCR", "true").lower() == "true"
    if not do_ocr:
        logger.info("OCR is disabled by configuration.")
        return ""
        
    ocr_engine = os.getenv("OCR_ENGINE", "auto").lower()
    
    # Try Surya first if engine is 'surya' or 'auto'
    if ocr_engine in ["surya", "auto"]:
        try:
            return parse_with_surya(images)
        except Exception:
            logger.info("Surya OCR failed or is not available. Falling back to VLM API OCR...")
            
    # Try VLM if engine is 'vlm' or fallback
    try:
        return parse_with_vlm(images)
    except Exception as e:
        logger.error(f"All OCR methods failed: {e}")
        return ""

def parse_document(file_path: str) -> str:
    """
    Main entry point for document parsing.
    Detects file extension and routes to the appropriate parser.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        text = parse_pdf_native(file_path)
        # If the extracted text is empty or too short (e.g. < 100 characters), it's likely a scanned PDF
        if len(text.strip()) < 100:
            logger.info("PDF has no native text layer or text is extremely short. Falling back to OCR.")
            return parse_image_ocr(file_path)
        return text
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_path)
    elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
        return parse_image_ocr(file_path)
    elif ext in [".txt", ".md"]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    else:
        logger.warning(f"Unsupported file type: {ext}. Attempting text reading.")
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading file: {e}")
            return ""
